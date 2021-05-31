#
# Copyright (c) nexB Inc. and others. All rights reserved.
# ScanCode is a trademark of nexB Inc.
# SPDX-License-Identifier: Apache-2.0
# See http://www.apache.org/licenses/LICENSE-2.0 for the license text.
# See https://github.com/nexB/scancode-toolkit for support or download.
# See https://aboutcode.org for more information about nexB OSS projects.
#
import docx
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import os

from commoncode.cliutils import PluggableCommandLineOption
from commoncode.cliutils import OUTPUT_GROUP
from plugincode.output import output_impl
from plugincode.output import OutputPlugin

from formattedcode import FileOptionType

# Tracing flags
TRACE = False


def logger_debug(*args):
    pass


if TRACE:
    import sys
    import logging

    logger = logging.getLogger(__name__)
    logging.basicConfig(stream=sys.stdout)
    logger.setLevel(logging.DEBUG)

    def logger_debug(*args):
        return logger.debug(' '.join(isinstance(a, str)
                                     and a or repr(a) for a in args))


@output_impl
class DocOutput(OutputPlugin):

    options = [
        PluggableCommandLineOption(('--doc',),
            type=FileOptionType(mode='wb', lazy=True),
            metavar='FILE',
            help='Write scan output as DOC to FILE.',
            help_group=OUTPUT_GROUP,
            sort_order=30),
    ]

    def is_enabled(self, doc, **kwargs):
        return doc

    def process_codebase(self, codebase, doc, **kwargs):
        results = self.get_files(codebase, **kwargs)
        write_doc(results, doc)


def write_doc(results, output_file):
    results = list(results)
    licensetext_dict = dict()

    rows = flatten_scan(results,licensetext_dict)
    
    """ Write the list data to a microsoft word document"""
    document = Document()
    licensetext_set = set()
    sections = document.sections
    section = sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_height = section.page_width
    section.page_width = section.page_height
    section.page_height = new_height
    table = document.add_table(rows=1, cols=5, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name of OSS Component'
    hdr_cells[1].text = 'Version of OSS Component'
    hdr_cells[2].text = 'Name and Version of License'
    hdr_cells[3].text = 'More Information'
    hdr_cells[4].text = 'Snippet URL'
    make_rows_bold(table.rows[0])
        
    componentdict = {}
    
    for scan_result_list in rows:
        for scan_list in scan_result_list:
            if isinstance(scan_list, dict):
                for key, val in scan_list.items():
                    if 'package_name' in scan_list.keys():
                        packagename = scan_list['package_name']
                        if key in ['license_name', 'copyright', 'package_version', 'url']:
                            tempdict = {}
                            if packagename in componentdict.keys():
                                    if key in (componentdict[packagename]).keys():
                                        mylist = []
                                        my_value = componentdict[packagename][key]
                                        mylist.append(my_value)
                                        mylist.extend([val])
                                        flatlist = []
                                        for elem in mylist:
                                            if type(elem) == list:
                                                for e in elem:
                                                    flatlist.append(e)
                                            else:
                                                flatlist.append(elem)
                                        flatlist = list(set(flatlist))
                                        componentdict[packagename][key] = flatlist
                                    else:
                                        tempdict[key] = val
                                        componentdict[packagename][key] = [val]
                            else:
                                tempdict[key] = [val]
                                componentdict[packagename] = tempdict   

    for k, v in componentdict.items():
        if k:
            row_cells = table.add_row().cells
            row_cells[0].text = k
            for k1, v1 in v.items():
                if 'package_version' in v.keys():
                    row_cells[1].text = v['package_version']
                if 'license_name' in v.keys():
                    strings = ','.join(filter(None, v['license_name']))
                    my_list = strings.split(",")
                    my_list = set(my_list)
                    licensetext_set.update(my_list)
                    my_string = '\n\n'.join(map(str, my_list)) 
                    row_cells[2].text = my_string
                    
                if 'copyright' in v.keys():
                    strings = ','.join(filter(None, v['copyright']))
                    my_list = strings.split(",")
                    if len(my_list) > 1:
                        my_list = set(my_list)
                        my_string = '\n\n'.join(map(str, my_list))
                        row_cells[3].text = my_string
                    else:
                        row_cells[3].text = my_list
                if 'url' in v.keys():
                    strings = ','.join(filter(None, v['url']))
                    my_list = strings.split(",")
                    my_list = set(my_list)
                    my_string = '\n'.join(map(str, my_list))
                    row_cells[4].text = my_string
    
    
    def add_hyperlink(paragraph, text, url):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    
        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
    
        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run ()
        r._r.append (hyperlink)
    
        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True
    
        return hyperlink

    phara = document.add_paragraph()
    phara.add_run("\n\nLink to license text of the found license:").bold = True
    for lname, ltext in licensetext_dict.items():
        phara.add_run("\n"+lname+" : ")
        add_hyperlink(phara, ltext, ltext)
    document.save(output_file)        
            
    
    
def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def flatten_scan(scan, licensetext_dict):
    """
    Flattening the sequence data in a single line-separated value and keying always by path,
    given a ScanCode `scan` results list.
    """
    
    """gets a list of orderedDict which holds all the details"""
    mainlist = []
    fg_val=0
    for scanned_file in scan:
        for k, v in scanned_file.items():
            if k == "path":
                newlist = dict()
                newlist['path'] = v
            elif k == "type":
                newlist['type'] = v
                if v=="directory":
                    fg_val+=1
                newlist['folder_group']=fg_val
            '''elif k == "folder_group":
                newlist['folder_group'] = v'''
            if type(v) is list:
                if k == "licenses":
                    for val1 in v:
                        if isinstance(val1, dict):
                            for keyname, keyvalue in val1.items():
                                keyvalue = str(keyvalue)
                                keyvalue = keyvalue.replace(',', '')
                                if keyname == "name":
                                    if 'license_name' in newlist.keys():
                                        if keyvalue not in newlist.get('license_name'):
                                            templist = newlist['license_name']
                                            keyvalue = ",".join([templist, keyvalue])
                                            newlist['license_name'] = keyvalue
                                    else:
                                        newlist['license_name'] = keyvalue
                                elif keyname == "text_url" and keyvalue != "":
                                    if 'license_text_url' in newlist.keys():
                                        if keyvalue not in newlist.get('license_text_url'):
                                            templist = newlist['license_text_url']
                                            keyvalue = ",".join([templist, keyvalue])
                                            newlist['license_text_url'] = keyvalue
                                            licensetext_dict[val1['name']] = templist
                                    else:
                                        newlist['license_text_url'] = keyvalue
                                        licensetext_dict[val1['name']] = keyvalue
                elif k == "copyrights":
                    for val1 in v:
                        if isinstance(val1, dict):
                            for keyname, keyvalue in val1.items():
                                if keyvalue is not None and isinstance(keyvalue, str):
                                    keyvalue = keyvalue.replace(",", "")
                                if keyname == "value":
                                    #print(keyvalue)
                                    if 'copyright' in newlist.keys():
                                        if keyvalue not in newlist.get('copyright'):
                                            templist = newlist['copyright']
                                            keyvalue = ",".join([templist, keyvalue])
                                            newlist['copyright'] = keyvalue
                                    else:
                                        newlist['copyright'] = keyvalue
                elif k == "packages":
                    for val1 in v:
                        if isinstance(val1, dict):
                            for keyname, keyvalue in val1.items():
                                if keyvalue is not None and isinstance(keyvalue, str):
                                    keyvalue = keyvalue.replace(",", "")
                                if keyname == "name":
                                    newlist['package_name'] = keyvalue
                                elif keyname == "version":  
                                    newlist['package_version'] = keyvalue
                                elif keyname == "homepage_url":
                                    newlist['package_homepage_url'] = keyvalue
                elif k == "urls":
                    for val1 in v:
                        if isinstance(val1, dict):
                            for keyname, keyvalue in val1.items():
                                if keyvalue is not None and isinstance(keyvalue, str):
                                    keyvalue = keyvalue.replace(",", "")
                                if keyname == "url":
                                    if 'url' in newlist.keys():
                                        if keyvalue not in newlist.get('url'):
                                            templist = newlist['url']
                                            keyvalue = ",".join([templist, keyvalue])
                                            newlist['url'] = keyvalue
                                    else:
                                        newlist['url'] = keyvalue
                              
        mainlist.append(newlist)

    previouspath = ''
    previouspackagename = ''
    previouspackageversion = ''
    previouspackageurl = ''
    flag = 0
    
    """get the previous path's package name and version"""
    for templist in mainlist: 
        if (templist['type'] == "directory") and ('package_name' not in templist.keys()) and (previouspath in templist['path']) and not templist['path'].endswith("node_modules"):
            if previouspackagename:
                templist['package_name'] = previouspackagename
                templist['package_version'] = previouspackageversion
                templist['package_homepage_url'] = previouspackageurl
                flag = 1
        else:
            flag = 0 
        if templist['type'] == "directory" and ('package_name' in templist.keys()) and flag == 0:
            previouspath = templist['path']
            previouspackagename = templist['package_name']
            previouspackageversion = templist['package_version']
            previouspackageurl = templist['package_homepage_url']
    
    """to print package name matching the folder group"""
    for sublist in mainlist:
        strippedpath, tail = os.path.split(sublist['path'])
        if (sublist['type'] == "directory") and ('package_name' not in sublist.keys()) and not sublist['path'].endswith("node_modules"):
            for templist in mainlist: 
                if templist['path'] == strippedpath and 'package_name' in templist.keys():
                    sublist['package_name'] = templist['package_name']
                    sublist['package_version'] = templist['package_version']
                    sublist['package_homepage_url'] = templist['package_homepage_url']
                    
        if 'package_name' in sublist.keys():
            fldr_grp = sublist['folder_group']
            for templist in mainlist:
                if templist['folder_group'] == fldr_grp and 'package_name' not in templist.keys():
                    templist['package_name'] = sublist['package_name']
                    templist['package_version'] = sublist['package_version']
                    templist['package_homepage_url'] = sublist['package_homepage_url']
    yield(mainlist)